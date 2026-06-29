
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.line_spacing = 1.3

def add_section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

def add_exp_header(doc, company, role, date):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(company + '  ')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(role + '  ')
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(0x2c, 0x6f, 0xaa)
    r3 = p.add_run(date)
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ===== HEADER =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run('吴 天 琦')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(1)
run = p.add_run('17717100070  |  17717100070@163.com  |  微信：wk923406')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
run = p.add_run('求职意向：供应链运营 / 数据分析  |  工作年限：3年')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Horizontal line
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(4)
pPr = p._p.get_or_add_pPr()
from docx.oxml.ns import qn
pBdr = pPr.makeelement(qn('w:pBdr'), {})
bottom = pBdr.makeelement(qn('w:bottom'), {qn('w:val'): 'single', qn('w:sz'): '6', qn('w:space'): '1', qn('w:color'): '2c6faa'})
pBdr.append(bottom)
pPr.append(pBdr)

# ===== SKILLS =====
add_section_title(doc, '专业技能')
skills = [
    ('供应链运营', '熟悉电商供应链售后链路，擅长通过数据分析定位缺货、履约等核心问题，联动商家与内部团队推动运营闭环。'),
    ('AI 自动化工作流', '熟练使用 Codex / Cursor 等 AI 编程工具辅助开发；掌握 Airflow 任务调度、Python 数据处理与 Jinja2 模板引擎，能独立搭建端到端数据自动化工作流。'),
    ('数据分析', '熟练 SQL 多表关联与窗口函数，精通 Excel 数据透视；能独立完成订单、库存、售后数据的取数、清洗与异常定位。'),
    ('自动化工具', '使用宜搭搭建业务流程表单与审批流；掌握 Python 脚本编写，熟悉 Playwright 自动化报告导出与定时任务配置。'),
    ('数据可视化', '使用 FBI、Tableau 搭建供应链监控看板；熟悉 ECharts 图表库与看板搭建。'),
    ('商家对接', '具备丰富的商家沟通经验，能有效协调多方诉求，推动售后规则优化与履约流程改善。'),
]
for name, desc in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run('▸ ' + name + '：')
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)

# ===== WORK EXPERIENCE =====
add_section_title(doc, '工作经历')

add_exp_header(doc, '阿里巴巴 · 速卖通', '供应链运营', '2025.10 - 至今')
bullets_ali = [
    '以缺货率为核心指标，通过 SQL 与看板工具监控 SKU 维度数据，主动识别库存不足、发货延迟等异常商家，定位售后履约瓶颈。',
    '运用 Codex 辅助搭建缺货率智能监控与自动预警工作流，数据清洗与报表效率提升 40%，问题发现效率提升 30%；同步搭建日报/周报 AI 自动生成系统，报告产出效率提升 50%+。',
    '参与售后体验优化，协同业务团队完善商家对接流程与异常处理机制，提升问题解决效率与商家配合度。',
]
for b in bullets_ali:
    add_bullet(doc, b)

add_exp_header(doc, '北京润合美品牌管理有限公司', '数据分析师', '2023.06 - 2025.07')
bullets_rh = [
    '负责蜂花、隆力奇等品牌天猫/京东店铺日常数据分析与报表自动化，通过 Power BI / Tableau 搭建运营看板，监控 GMV、转化率等核心指标。',
    '使用 Codex 辅助编写 Python 脚本，实现推广数据自动抓取与整理，对账效率提升 60%。',
    '搭建宜搭数据上报流程，统一各渠道销售数据入口，减少跨平台手工对账时间。',
]
for b in bullets_rh:
    add_bullet(doc, b)

# ===== AI WORKFLOW PROJECTS =====
add_section_title(doc, 'AI 自动化工作流项目')

projects = [
    ('项目一：SKU 缺货率智能监控与自动预警', [
        '每日自动监控 5,000+ SKU，通过 SQL + Python 自动化采集数据，异常检测模型实时识别缺货风险并分级（高/中/低）预警。',
        'Codex 辅助编写 10+ SQL 模板与 3 个 Python 脚本（数据拉取、异常检测、企微推送），搭建 FBI 监控看板 + Crontab 定时调度。',
        '问题发现效率提升 30%，缺货率下降约 15%，人工巡检从 2h/天降至全自动。',
        '实体产出：10+ SQL 模板、800 行 Python 脚本、FBI 看板配置文件、预警消息样例。',
    ]),
    ('项目二：供应链日报/周报 AI 自动生成', [
        '搭建「数据采集 → AI 分析注释 → Jinja2 渲染 → 邮件推送」全链路自动化，覆盖 GMV、缺货率、售后率等 8+ 核心指标。',
        'Codex 辅助编写报告引擎（~200 行 Python），设计日报/周报 HTML 模板，自动截图 FBI 图表嵌入报告。',
        '报告产出效率提升 50%+，从手工 1.5h/天 → 全自动 5 分钟，支持邮件/企微双通道推送。',
        '实体产出：报告引擎脚本、日报/周报 Jinja2 模板、配置 YAML、报告产出样例（HTML 可直接浏览器打开）。',
    ]),
    ('项目三：商家活动库存编辑审批自动化', [
        '基于宜搭低代码平台搭建三级审批流，根据库存变更量自动分级路由（<500 件 L1 / 500-2000 L2 / >2000 L3）。',
        'Codex 辅助生成宜搭表单 Schema（含校验规则）、分级路由 JS 扩展逻辑、API 回写脚本（含 HMAC 签名与重试机制）。',
        '审批时效从人工 4~8h → 系统 <1h，100% 审批留痕可追溯，人工审批遗漏风险降至 0。',
        '实体产出：宜搭表单 + 三级审批流配置、API 回写脚本（~120 行）、四种场景通知模板。',
    ]),
]

for proj_name, proj_bullets in projects:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(proj_name)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
    for b in proj_bullets:
        add_bullet(doc, b)

# ===== OTHER PROJECTS =====
add_section_title(doc, '其他项目经历')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
r = p.add_run('蜂花洗护品牌年度库存与销售分析（2024）')
r.bold = True
r.font.size = Pt(10.5)
for b in [
    '汇总经营、库存、推广等 20+ 核心指标，完成同比与季度拆解，支撑库存优化决策。',
    '使用宜搭搭建库存数据上报与审批流程，实现各渠道库存数据实时同步，减少数据录入错误。',
    '分析渠道 ROI 并优化投放结构，全年推广费占比由 8.3% 降至 7.04%，GMV 达 6572 万（超 KPI 6%）。',
]:
    add_bullet(doc, b)

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
r = p.add_run('蜂花老客 RFM 精细化运营（2024）')
r.bold = True
r.font.size = Pt(10.5)
for b in [
    '清洗 66.6 万老客数据，运用 K-means 聚类算法划分客群，支撑精准补货与库存规划。',
    '利用 AI 工具辅助客群分析与策略生成，按季度落地差异化运营策略。',
    '老客复购率由 15% 提升至 19%，新增 GMV 172.9 万。',
]:
    add_bullet(doc, b)

# ===== EDUCATION =====
add_section_title(doc, '教育经历')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
r1 = p.add_run('西安邮电大学（非全日制）  ')
r1.bold = True
r1.font.size = Pt(10)
r2 = p.add_run('计算机科学与技术 · 本科')
r2.font.size = Pt(10)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
r1 = p.add_run('石家庄信息工程职业学院  ')
r1.bold = True
r1.font.size = Pt(10)
r2 = p.add_run('艺术设计 · 大专')
r2.font.size = Pt(10)

# ===== PORTFOLIO LINK =====
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(4)
r = p.add_run('📂 AI 工作流作品集')
r.bold = True
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('https://2107928138-del.github.io/ai-workflow-portfolio/landing.html')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x2c, 0x6f, 0xaa)
r.underline = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('（面试时可扫码或直接打开链接查看完整项目源码与产出样例）')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

doc.save('/Users/wutianqi/Documents/Codex/2026-06-29/ni-h/outputs/吴天琦_供应链运营_简历.docx')
print('DOCX saved')
