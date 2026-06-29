
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.line_spacing = 1.3

def section_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

def bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    for r in p.runs:
        r.font.size = Pt(10)

def exp_header(doc, company, role, date):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(company + '  ')
    r1.bold = True
    r1.font.size = Pt(11)
    r2 = p.add_run(role + '  ')
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = RGBColor(0x2c, 0x6f, 0xaa)
    r3 = p.add_run(date)
    r3.font.size = Pt(10)
    r3.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

def proj_title(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

# ===== HEADER =====
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
run = p.add_run('吴 天 琦')
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(1)
run = p.add_run('17717100070  |  17717100070@163.com  |  微信：wk923406')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
run = p.add_run('求职意向：AI 应用工程师 / AI 产品运营 / 智能运营  |  3年经验')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# horizontal line
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(0)
p.paragraph_format.space_after = Pt(4)
pPr = p._p.get_or_add_pPr()
pBdr = pPr.makeelement(qn('w:pBdr'), {})
bottom = pBdr.makeelement(qn('w:bottom'), {qn('w:val'): 'single', qn('w:sz'): '6', qn('w:space'): '1', qn('w:color'): '2c6faa'})
pBdr.append(bottom)
pPr.append(pBdr)

# ===== SKILLS =====
section_title(doc, '核心能力')

skills = [
    ('🤖 AI 工具链 & 应用开发', '熟练使用 Codex / Cursor 等 AI IDE 进行 Prompt 工程与自动化脚本开发；能独立将业务需求转化为 AI 驱动的自动化方案，具备 3 套完整 AI 工作流系统从 0 到 1 的搭建经验。'),
    ('⚙️ 自动化工作流搭建', '掌握 Airflow 任务调度、Python 数据处理、Jinja2 模板引擎，能搭建「数据采集 → AI 分析 → 模板渲染 → 自动推送」端到端自动化管线。'),
    ('🔗 低代码 + API 集成', '使用宜搭搭建业务流程表单与多级审批流；熟悉 REST API 对接、HMAC 签名鉴权、回调机制，实现低代码平台与业务系统双向打通。'),
    ('📊 数据分析 & 可视化', '精通 SQL 多表关联与窗口函数；熟练使用 FBI、Tableau 搭建监控看板；掌握 ECharts 前端图表库，能独立完成数据从采集到可视化的全链路。'),
    ('💡 业务驱动 AI 落地', '具备供应链、电商真实业务场景经验，能快速理解业务痛点并转化为可落地的 AI 自动化方案，产出有量化效果（效率提升 30%-50%+）。'),
]
for name, desc in skills:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run('▸ ' + name + '：')
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = p.add_run(desc)
    r2.font.size = Pt(10)

# ===== WORK EXPERIENCE =====
section_title(doc, '工作经历')

exp_header(doc, '阿里巴巴 · 速卖通', 'AI 自动化运营（供应链方向）', '2025.10 - 至今')
for b in [
    '以 AI 编程工具（Codex）为核心生产力工具，独立搭建三套 AI 自动化工作流系统——SKU 缺货率智能监控预警、供应链日报/周报 AI 自动生成、商家库存编辑审批自动化，实现从「需求分析 → AI 辅助编码 → 测试部署 → 定时调度」的全生命周期交付。',
    '通过 SQL + Python + Airflow 构建数据自动化采集与异常检测管线，问题发现效率提升 30%，数据清洗与报表效率提升 40%，人工巡检时间从 2 小时/天 → 全自动零人工。',
    '利用 Jinja2 模板引擎 + SMTP 搭建日报/周报 AI 自动生成与推送系统，覆盖 8+ 核心业务指标，报告产出效率提升 50%+。',
    '基于宜搭低代码平台搭建三级审批流，通过 API 集成实现审批通过后自动回写库存系统，审批时效从人工 4~8 小时 → 系统 1 小时以内。',
]:
    bullet(doc, b)

exp_header(doc, '北京润合美品牌管理有限公司', '数据分析师', '2023.06 - 2025.07')
for b in [
    '负责蜂花、隆力奇等品牌天猫/京东店铺数据分析与报表自动化，通过 Power BI / Tableau 搭建运营看板。',
    '使用 Codex 辅助编写 Python 脚本实现推广数据自动抓取与对账，效率提升 60%；利用宜搭搭建数据上报流程，统一多平台数据入口。',
]:
    bullet(doc, b)

# ===== AI PROJECTS (MAIN SHOWCASE) =====
section_title(doc, 'AI 自动化工作流项目（核心亮点）')

projects = [
    ('项目一：SKU 缺货率智能监控与自动预警系统', [
        '场景：速卖通平台 5,000+ SKU 的缺货率日常监控，需及时发现异常商家并预警。',
        'AI 参与：Codex 辅助编写全部代码（SQL 模板库 + 3 个 Python 脚本 + FBI 看板配置），Prompt 工程驱动开发，开发周期缩短 60%+。',
        '技术栈：Python + SQL + Pandas + Codex + FBI + Cron + 企业微信 API。',
        '成果：问题发现效率 ↑30%，缺货率 ↓~15%，人工巡检从 2h/天 → 全自动。实体产出：10+ SQL 模板、800 行 Python 脚本、预警消息样例。',
    ]),
    ('项目二：供应链日报/周报 AI 自动生成系统', [
        '场景：运营团队每日/每周需手动整理数据、截图、写分析、发邮件，耗时 1.5h+/天。',
        'AI 参与：Codex 辅助编写报告引擎（~200 行 Python）+ Jinja2 模板设计 + FBI 图表自动截图逻辑。AI 生成初步分析注释，人工审核后推送。',
        '技术栈：Python + Jinja2 + Codex + SMTP + Cron + ECharts。',
        '成果：报告产出效率 ↑50%+，从手工 1.5h/天 → 全自动 5 分钟。实体产出：日报/周报 HTML 模板、产出样例（可直接浏览器打开）。',
    ]),
    ('项目三：商家库存编辑审批自动化系统', [
        '场景：大促期间商家频繁提交库存编辑申请，人工审批耗时 4~8h，易遗漏。',
        'AI 参与：Codex 辅助生成宜搭表单 JSON Schema（含校验规则）、分级路由 JS 逻辑、API 回写脚本（含 HMAC 签名鉴权 + 重试机制）。',
        '技术栈：宜搭 + Codex + REST API + 钉钉/企微通知 + 审计日志。',
        '成果：审批时效 4~8h → <1h，100% 留痕可追溯，0 遗漏风险。实体产出：表单 Schema、审批流配置、回写脚本（~120 行）、通知模板。',
    ]),
]

for proj_name, proj_bullets in projects:
    proj_title(doc, proj_name)
    for b in proj_bullets:
        bullet(doc, b)

# ===== EDUCATION =====
section_title(doc, '教育经历')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
r1 = p.add_run('西安邮电大学（非全日制）  ')
r1.bold = True
r1.font.size = Pt(10)
r2 = p.add_run('计算机科学与技术 · 本科')
r2.font.size = Pt(10)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
r1 = p.add_run('石家庄信息工程职业学院  ')
r1.bold = True
r1.font.size = Pt(10)
r2 = p.add_run('艺术设计 · 大专')
r2.font.size = Pt(10)

# ===== SELF-ASSESSMENT =====
section_title(doc, '自我评价')

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run('具备「业务理解 + AI 工具链 + 工程落地」三位一体的复合能力——能独立将模糊的业务需求拆解为可执行的 AI 自动化方案，并用 Codex 等 AI IDE 高效完成编码与部署。')
r.font.size = Pt(10)
p2 = doc.add_paragraph()
r2 = p2.add_run('3 套完整 AI 工作流系统从 0 到 1 的搭建经验，覆盖数据采集、异常检测、报表生成、审批自动化等典型业务场景，每套系统均有量化的效率提升数据。')
r2.font.size = Pt(10)

# ===== PORTFOLIO =====
section_title(doc, 'AI 工作流作品集')

img_dir = os.path.dirname(os.path.abspath(__file__))
images = [
    ('screenshot_p1.png', 'SKU 缺货率智能监控系统'),
    ('screenshot_p2.png', '日报/周报 AI 自动生成系统'),
    ('screenshot_p3.png', '商家库存编辑审批自动化系统'),
]
for fname, caption in images:
    img_path = img_dir + '/' + fname
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(caption)
        r.font.size = Pt(9)
        r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(4)
r = p.add_run('📂 完整作品集（源码 + 配置 + 产出样例）')
r.bold = True
r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('https://2107928138-del.github.io/ai-workflow-portfolio/landing.html')
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x2c, 0x6f, 0xaa)
r.underline = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('（面试时可扫码或直接打开链接，查看完整项目源码与交互界面）')
r.font.size = Pt(9)
r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

out = '/Users/wutianqi/Documents/Codex/2026-06-29/ni-h/outputs/吴天琦_AI应用_简历.docx'
doc.save(out)
print('OK: ' + out)
